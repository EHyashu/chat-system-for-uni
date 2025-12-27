from dataclasses import dataclass
from pathlib import Path
from typing import List

from docx import Document as DocxDocument
from pypdf import PdfReader


@dataclass
class RawDocument:
    text: str
    source: str


def load_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def load_pdf(path: Path) -> str:
    reader = PdfReader(path)
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def load_docx(path: Path) -> str:
    doc = DocxDocument(path)
    return "\n".join(p.text for p in doc.paragraphs)


def load_documents(documents_dir: Path) -> List[RawDocument]:
    docs: List[RawDocument] = []
    for path in documents_dir.rglob("*"):
        if not path.is_file():
            continue
        ext = path.suffix.lower()
        text = ""
        if ext == ".txt":
            text = load_txt(path)
        elif ext == ".pdf":
            text = load_pdf(path)
        elif ext == ".docx":
            text = load_docx(path)
        else:
            continue

        text = text.strip()
        if not text:
            continue

        docs.append(RawDocument(text=text, source=path.name))
    return docs
